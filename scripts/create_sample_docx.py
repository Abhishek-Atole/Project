from docx import Document

def create_sample_docx(file_path):
    """
    Create a sample Word document with questions and choices.
    """
    doc = Document()
    doc.add_heading("Sample MCQ Questions", level=1)

    # Add a question with choices
    doc.add_paragraph("What is 2 + 2?")
    doc.add_paragraph("A. 3")
    doc.add_paragraph("B. 4")
    doc.add_paragraph("C. 5")
    doc.add_paragraph("D. 6")

    # Add another question
    doc.add_paragraph("\nSolve the equation: x^2 - 4 = 0")
    doc.add_paragraph("A. x = 2")
    doc.add_paragraph("B. x = -2")
    doc.add_paragraph("C. x = ±2")
    doc.add_paragraph("D. None")

    # Save the document
    doc.save(file_path)
    print(f"Sample Word document created at: {file_path}")

# Generate the file in the input_files directory
create_sample_docx("../input_files/sample.docx")